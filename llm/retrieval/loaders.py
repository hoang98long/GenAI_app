# llm/retrieval/loaders.py

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

from loguru import logger
from pypdf import PdfReader
import docx


@dataclass
class RawDocument:
    text: str
    metadata: Dict[str, object]


def load_pdf(path: Path) -> List[RawDocument]:
    reader = PdfReader(str(path))
    docs: List[RawDocument] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        docs.append(
            RawDocument(
                text=text,
                metadata={
                    "source": str(path),
                    "page": i + 1,
                    "type": "pdf",
                },
            )
        )
    logger.info("Loaded {} pages from PDF {}", len(docs), path)
    return docs


def load_docx(path: Path) -> List[RawDocument]:
    document = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    docs: List[RawDocument] = []
    for i, para in enumerate(paragraphs):
        docs.append(
            RawDocument(
                text=para,
                metadata={
                    "source": str(path),
                    "paragraph": i + 1,
                    "type": "docx",
                },
            )
        )
    logger.info("Loaded {} paragraphs from DOCX {}", len(docs), path)
    return docs


def load_txt(path: Path, encoding: str = "utf-8") -> List[RawDocument]:
    text = path.read_text(encoding=encoding, errors="ignore").strip()
    if not text:
        return []
    # Có thể tách theo dòng hoặc đoạn, ở đây mình giữ nguyên 1 block
    return [
        RawDocument(
            text=text,
            metadata={
                "source": str(path),
                "type": "txt",
            },
        )
    ]


def load_file(path: Path) -> List[RawDocument]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix in (".docx", ".doc"):
        return load_docx(path)
    if suffix in (".txt", ".md"):
        return load_txt(path)

    logger.warning("Unsupported file type for {}. Skipping.", path)
    return []
